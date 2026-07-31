"""Exports every Markdown file in docs/ to PDF, Word (.docx) and HTML.

Run with:
    python scripts/export_documentation.py

Produces docs/exports/<name>.pdf, docs/exports/<name>.docx and
docs/exports/<name>.html for each docs/*.md file — the academic deliverables
requested for the thesis (rapports exportables PDF/Word/HTML).

The Markdown parsing here is intentionally minimal (headings, paragraphs,
bullet lists, pipe tables): it only needs to handle the specific formatting
used in this project's own docs/*.md files, not arbitrary Markdown.
"""

import re
from pathlib import Path

import markdown as md_lib
from docx import Document
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

DOCS_DIR = Path(__file__).resolve().parent.parent / "docs"
EXPORTS_DIR = DOCS_DIR / "exports"


def _parse_markdown_blocks(text: str) -> list[dict]:
    """Split a markdown document into typed blocks: heading/paragraph/bullets/table."""

    blocks: list[dict] = []
    lines = text.splitlines()
    i = 0
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        if paragraph_buffer:
            blocks.append({"type": "paragraph", "text": " ".join(paragraph_buffer).strip()})
            paragraph_buffer.clear()

    while i < len(lines):
        line = lines[i]

        if line.startswith("# "):
            flush_paragraph()
            blocks.append({"type": "h1", "text": line[2:].strip()})
        elif line.startswith("## "):
            flush_paragraph()
            blocks.append({"type": "h2", "text": line[3:].strip()})
        elif line.startswith("### "):
            flush_paragraph()
            blocks.append({"type": "h3", "text": line[4:].strip()})
        elif line.strip().startswith("|"):
            flush_paragraph()
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = [
                [cell.strip() for cell in row.strip().strip("|").split("|")]
                for row in table_lines
                if not re.match(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)*\|?$", row.strip())
            ]
            blocks.append({"type": "table", "rows": rows})
            continue
        elif line.strip().startswith("- "):
            flush_paragraph()
            bullets = []
            while i < len(lines) and lines[i].strip().startswith("- "):
                bullets.append(lines[i].strip()[2:])
                i += 1
            blocks.append({"type": "bullets", "items": bullets})
            continue
        elif line.strip() == "":
            flush_paragraph()
        else:
            paragraph_buffer.append(line.strip())

        i += 1

    flush_paragraph()
    return blocks


def _strip_markdown_inline(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


def export_pdf(name: str, blocks: list[dict], output_path: Path) -> None:
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(str(output_path), pagesize=A4, title=name)
    story: list = []

    for block in blocks:
        if block["type"] == "h1":
            story.append(Paragraph(_strip_markdown_inline(block["text"]), styles["Title"]))
            story.append(Spacer(1, 12))
        elif block["type"] == "h2":
            story.append(Spacer(1, 8))
            story.append(Paragraph(_strip_markdown_inline(block["text"]), styles["Heading2"]))
        elif block["type"] == "h3":
            story.append(Paragraph(_strip_markdown_inline(block["text"]), styles["Heading3"]))
        elif block["type"] == "paragraph":
            story.append(Paragraph(_strip_markdown_inline(block["text"]), styles["Normal"]))
            story.append(Spacer(1, 6))
        elif block["type"] == "bullets":
            items = [ListItem(Paragraph(_strip_markdown_inline(b), styles["Normal"])) for b in block["items"]]
            story.append(ListFlowable(items, bulletType="bullet"))
            story.append(Spacer(1, 6))
        elif block["type"] == "table" and block["rows"]:
            table = Table(block["rows"])
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2563eb")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                        ("FONTSIZE", (0, 0), (-1, -1), 7),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 10))

    doc.build(story)


def export_docx(name: str, blocks: list[dict], output_path: Path) -> None:
    document = Document()
    document.styles["Normal"].font.size = Pt(10.5)

    for block in blocks:
        if block["type"] == "h1":
            document.add_heading(_strip_markdown_inline(block["text"]), level=0)
        elif block["type"] == "h2":
            document.add_heading(_strip_markdown_inline(block["text"]), level=1)
        elif block["type"] == "h3":
            document.add_heading(_strip_markdown_inline(block["text"]), level=2)
        elif block["type"] == "paragraph":
            document.add_paragraph(_strip_markdown_inline(block["text"]))
        elif block["type"] == "bullets":
            for item in block["items"]:
                document.add_paragraph(_strip_markdown_inline(item), style="List Bullet")
        elif block["type"] == "table" and block["rows"]:
            rows = block["rows"]
            table = document.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Light Grid Accent 1"
            for r, row in enumerate(rows):
                for c, cell in enumerate(row):
                    table.cell(r, c).text = _strip_markdown_inline(cell)

    document.save(str(output_path))


def export_html(name: str, raw_markdown: str, output_path: Path) -> None:
    body = md_lib.markdown(raw_markdown, extensions=["tables"])
    html = f"""<!doctype html>
<html lang="fr"><head><meta charset="utf-8"><title>{name}</title>
<style>
body {{ font-family: Segoe UI, Arial, sans-serif; max-width: 900px; margin: 2rem auto; padding: 0 1rem; color: #1e293b; }}
h1, h2, h3 {{ color: #2563eb; }}
table {{ border-collapse: collapse; width: 100%; margin: 1rem 0; }}
th, td {{ border: 1px solid #cbd5e1; padding: 6px 10px; font-size: 0.9rem; }}
th {{ background: #2563eb; color: white; }}
</style></head><body>
{body}
</body></html>"""
    output_path.write_text(html, encoding="utf-8")


def main() -> None:
    EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
    md_files = sorted(DOCS_DIR.glob("*.md"))

    for md_path in md_files:
        name = md_path.stem
        raw_text = md_path.read_text(encoding="utf-8")
        blocks = _parse_markdown_blocks(raw_text)

        export_pdf(name, blocks, EXPORTS_DIR / f"{name}.pdf")
        export_docx(name, blocks, EXPORTS_DIR / f"{name}.docx")
        export_html(name, raw_text, EXPORTS_DIR / f"{name}.html")

        print(f"Exported {name}: PDF, DOCX, HTML")


if __name__ == "__main__":
    main()
